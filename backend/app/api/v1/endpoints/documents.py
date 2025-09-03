import os
import uuid
from typing import List
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.auth import get_current_user
from app.core.config import settings
from app.models.document import Document
from app.models.analysis import Analysis
from app.models.user import User
from app.schemas.document import DocumentResponse, DocumentCreate
from app.schemas.analysis import AnalysisResponse
from app.services.rag_service import RAGService
import PyPDF2
from docx import Document as DocxDocument

router = APIRouter()
rag_service = RAGService()

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a new document"""
    # Validate file type
    if not file.filename or "." not in file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid filename"
        )
    
    file_extension = file.filename.split(".")[-1].lower()
    if file_extension not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type {file_extension} not allowed"
        )
    
    # Validate file size
    if file.size > settings.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large"
        )
    
    # Generate unique filename
    unique_filename = f"{uuid.uuid4()}.{file_extension}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
    
    # Save file
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Get user
    user = db.query(User).filter(User.email == current_user.email).first()
    
    # Create document record
    db_document = Document(
        user_id=user.id,
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=file.size,
        mime_type=file.content_type or "application/octet-stream",
        status="uploaded"
    )
    
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    return db_document

@router.get("/", response_model=List[DocumentResponse])
async def get_documents(
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all documents for current user"""
    user = db.query(User).filter(User.email == current_user.email).first()
    documents = db.query(Document).filter(Document.user_id == user.id).all()
    return documents

@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get specific document by ID"""
    user = db.query(User).filter(User.email == current_user.email).first()
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return document

@router.post("/{document_id}/analyze", response_model=AnalysisResponse)
async def analyze_document(
    document_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Analyze document using RAG"""
    user = db.query(User).filter(User.email == current_user.email).first()
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Extract text from document
    text = ""
    try:
        if document.mime_type == "application/pdf":
            with open(document.file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        elif document.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(document.file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            with open(document.file_path, "r", encoding="utf-8") as file:
                text = file.read()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error reading document: {str(e)}"
        )
    
    # Process with RAG
    rag_result = rag_service.process_document(text, str(document_id))
    
    # Create analysis record
    analysis = Analysis(
        document_id=document_id,
        analysis_type="summary",
        original_text=text[:1000],  # Store first 1000 chars
        simplified_text=rag_result["analysis"],
        confidence_score=rag_result["confidence_score"],
        processing_time=rag_result["processing_time"]
    )
    
    db.add(analysis)
    
    # Update document status
    document.status = "analyzed"
    
    db.commit()
    db.refresh(analysis)
    
    return analysis

@router.get("/{document_id}/analysis", response_model=List[AnalysisResponse])
async def get_document_analyses(
    document_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all analyses for a document"""
    user = db.query(User).filter(User.email == current_user.email).first()
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    analyses = db.query(Analysis).filter(Analysis.document_id == document_id).all()
    return analyses
